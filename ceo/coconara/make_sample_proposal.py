"""サンプル提案書を生成（APIなし、ハードコード版）"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, datetime

OUT = os.path.join(os.path.dirname(__file__), "samples")
os.makedirs(OUT, exist_ok=True)

NAVY  = RGBColor(0x1A, 0x3A, 0x6E)
AQUA  = RGBColor(0x00, 0xB4, 0xD8)
GOLD  = RGBColor(0xD4, 0xAF, 0x37)
GRAY  = RGBColor(0x80, 0x90, 0xA0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def make_proposal():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3)
        sec.right_margin = Cm(2.5)

    # ── 表紙 ─────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover_title.add_run("AI業務効率化クラウドサービス")
    r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = NAVY

    cover_sub = doc.add_paragraph()
    cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = cover_sub.add_run("導入提案書")
    r2.font.size = Pt(22); r2.font.bold = True; r2.font.color.rgb = NAVY

    doc.add_paragraph()

    for label, val in [("提案先", "株式会社サンプル商事  御中"),
                       ("日　付", datetime.date.today().strftime("%Y年%m月%d日")),
                       ("提案者", "AIエージェント事業部")]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_l = p.add_run(f"{label}："); r_l.font.size = Pt(13); r_l.font.color.rgb = GRAY
        r_v = p.add_run(val); r_v.font.size = Pt(13); r_v.font.bold = True

    doc.add_page_break()

    # ── 共通ヘッダー関数 ──────────────────────────────────────
    def add_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after  = Pt(6)
        # 左ボーダー風 (下線で代用)
        r = p.add_run(text)
        r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = NAVY
        # 下線
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(8)
        r2 = p2.add_run("━" * 30)
        r2.font.size = Pt(10); r2.font.color.rgb = AQUA

    def add_body(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.size = Pt(11)

    def add_bullet(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run("■ " + text)
        r.font.size = Pt(11)

    # ── 1. エグゼクティブサマリー ─────────────────────────────
    add_heading("1. エグゼクティブサマリー")
    add_body(
        "本提案は、貴社の業務プロセスに最新AI技術を導入することで、年間200時間以上の業務削減と"
        "コスト30%削減を実現するものです。競合他社がExcel手作業に頼る中、AIクラウドサービスを"
        "活用することで圧倒的な競争優位を確立し、意思決定の高速化・品質均一化を同時に達成します。"
    )

    kv = [("期待効果", "年間コスト削減 約180万円 / ROI 340%（初年度）"),
          ("導入期間", "最短2週間〜本格稼働"),
          ("リスク",   "初月無料トライアルにより導入リスクゼロ")]
    for k, v in kv:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r1 = p.add_run(f"【{k}】"); r1.font.bold=True; r1.font.color.rgb=NAVY; r1.font.size=Pt(11)
        r2 = p.add_run(f"  {v}"); r2.font.size=Pt(11)

    # ── 2. 現状の課題 ─────────────────────────────────────────
    add_heading("2. 現状の課題")
    bullets = [
        "手作業データ入力に月平均40時間を消費（ミス率3%）",
        "レポート作成に毎月3営業日を費やし、意思決定が最大1週間遅延",
        "属人化した業務フローにより、担当者不在時に業務が完全停止",
    ]
    for b in bullets: add_bullet(b)

    # ── 3. 提案内容 ───────────────────────────────────────────
    add_heading("3. 提案内容")
    funcs = [
        ("自動データ処理", "OCR＋AIで書類を自動読み取り・入力。手入力作業を87%削減。"),
        ("レポート自動生成", "数値を入れるだけで経営レポートを即時生成。3日→2時間に短縮。"),
        ("業務フロー標準化", "AIがステップバイステップでガイド。誰でも同品質で実行可能に。"),
    ]
    for title, desc in funcs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"● {title}  "); r1.font.bold=True; r1.font.size=Pt(11); r1.font.color.rgb=NAVY
        r2 = p.add_run(desc); r2.font.size=Pt(11)

    # ── 4. 導入効果・ROI ──────────────────────────────────────
    add_heading("4. 導入効果・ROI")
    rows = [
        ("項目", "現在", "導入後", "改善率"),
        ("データ入力工数", "40時間/月", "5時間/月", "87%削減"),
        ("レポート作成",   "3日/月",    "2時間/月", "93%削減"),
        ("ミス発生率",     "3%",        "0.1%以下", "97%削減"),
        ("年間コスト削減", "―",         "約180万円",  "―"),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    table.style = "Table Grid"
    col_w = [Cm(5), Cm(3.5), Cm(3.5), Cm(3)]
    for i, row in enumerate(rows):
        for j, (cell_text, cw) in enumerate(zip(row, col_w)):
            cell = table.cell(i, j)
            cell.width = cw
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(cell_text)
            r.font.size = Pt(11)
            if i == 0:
                r.font.bold = True
                r.font.color.rgb = WHITE
                set_cell_bg(cell, "1A3A6E")
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("初年度投資回収率（ROI）: 340%"); r.font.bold=True; r.font.size=Pt(12); r.font.color.rgb=NAVY

    # ── 5. 導入スケジュール ───────────────────────────────────
    add_heading("5. 導入スケジュール")
    phases = [
        ("フェーズ1（1〜2週間）", "環境構築・初期設定・データ移行"),
        ("フェーズ2（3〜4週間）", "操作研修・並行運用・品質確認"),
        ("フェーズ3（2ヶ月目〜）","本格稼働・効果測定・継続改善"),
    ]
    for ph, desc in phases: add_bullet(f"{ph}  {desc}")

    # ── 6. 料金プラン ─────────────────────────────────────────
    add_heading("6. 料金プラン")
    plans = [
        ("スタンダード", "月額 ¥50,000", "〜10ユーザー"),
        ("プロ",         "月額 ¥100,000","〜50ユーザー・API連携込み"),
        ("エンタープライズ","個別見積もり","無制限・専任サポート"),
    ]
    for name, price, note in plans:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r1 = p.add_run(f"【{name}】 "); r1.font.bold=True; r1.font.size=Pt(11); r1.font.color.rgb=NAVY
        r2 = p.add_run(f"{price}  "); r2.font.bold=True; r2.font.size=Pt(11)
        r3 = p.add_run(f"（{note}）"); r3.font.size=Pt(10); r3.font.color.rgb=GRAY
    doc.add_paragraph()
    note = doc.add_paragraph()
    rn = note.add_run("★ 初月無料トライアル実施中（いつでもキャンセル可）")
    rn.font.bold=True; rn.font.size=Pt(11); rn.font.color.rgb=GOLD

    # ── 7. よくある質問 ───────────────────────────────────────
    add_heading("7. よくある質問")
    faqs = [
        ("既存システムとの連携は？", "Excel・kintone・Salesforceなど主要システムとのAPI連携に対応しています。"),
        ("セキュリティは？",         "国内サーバー管理、通信暗号化、アクセスログ管理を徹底しています。"),
        ("操作が難しくないですか？", "専任CSチームが導入から定着まで伴走します。平均習得時間は2時間です。"),
    ]
    for q, a in faqs:
        pq = doc.add_paragraph()
        pq.paragraph_format.left_indent = Cm(0.5)
        pq.paragraph_format.space_after = Pt(2)
        rq = pq.add_run(f"Q. {q}"); rq.font.bold=True; rq.font.size=Pt(11); rq.font.color.rgb=NAVY
        pa = doc.add_paragraph()
        pa.paragraph_format.left_indent = Cm(1)
        pa.paragraph_format.space_after = Pt(8)
        ra = pa.add_run(f"A. {a}"); ra.font.size=Pt(11)

    # ── 8. 会社概要 ───────────────────────────────────────────
    add_heading("8. 会社概要")
    infos = [
        ("名称", "AIエージェント事業部"),
        ("設立", "2026年"),
        ("所在地", "東京都"),
        ("使用AI", "Anthropic Claude（世界最高水準モデル）"),
        ("実績", "中小企業を中心に多数の導入実績"),
    ]
    for k, v in infos:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r1 = p.add_run(f"{k}："); r1.font.bold=True; r1.font.size=Pt(11)
        r2 = p.add_run(v); r2.font.size=Pt(11)

    out_path = os.path.join(OUT, "sample_proposal_AI業務効率化.docx")
    doc.save(out_path)
    print(f"✅ サンプル提案書を保存: {out_path}")

if __name__ == "__main__":
    make_proposal()
