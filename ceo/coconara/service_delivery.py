"""
AI提案書・営業資料 自動生成システム
cocoanara 出品サービス「激速納品」の納品エンジン

使い方:
    python service_delivery.py --client "株式会社XXX" --service "クラウド会計" --audience "中小企業経営者"

出力: samples/<client>_proposal.docx + .pdf

APIキー設定:
    環境変数 ANTHROPIC_API_KEY か、このスクリプトと同じディレクトリの .env ファイル
"""
import anthropic
import argparse
import datetime
import os
from pathlib import Path

# .env ファイルから APIキーを読み込む（python-dotenv がなくても動く最小実装）
def _load_dotenv():
    env_path = Path(__file__).parent / ".env"
    if not env_path.exists():
        env_path = Path(__file__).parent.parent.parent / ".env"
    if env_path.exists():
        for line in env_path.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, _, v = line.partition("=")
                os.environ.setdefault(k.strip(), v.strip())

_load_dotenv()
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Claude API ----
client = anthropic.Anthropic()

SYSTEM_PROMPT = """あなたはトップクラスのビジネスコンサルタントです。
依頼された条件に基づき、クライアントに刺さる高品質なビジネス提案書を作成してください。
- 数字・実績・ROIを必ず含める
- 読み手が「これは自分ごとだ」と感じる表現を使う
- セクションは明確に区切り、スキャナビリティを高める
- 日本語ビジネス文書として自然な文体で"""

def generate_proposal_text(client_name: str, service: str, audience: str, budget: str, extra: str) -> dict:
    """Claude APIで提案書テキストを生成"""
    prompt = f"""
以下の条件で、A4・8〜10ページ相当のビジネス提案書を作成してください。

【クライアント企業名】{client_name}
【提案サービス/製品】{service}
【ターゲット】{audience}
【予算感】{budget}
【追加情報】{extra if extra else "なし"}

以下のセクションを必ず含めてください:
1. エグゼクティブサマリー（300字）
2. 課題・背景（現状の問題を3点）
3. 提案内容（サービス詳細・特徴・強み）
4. 導入効果・ROI（数値付き）
5. 導入スケジュール（フェーズ別）
6. 料金・プラン
7. よくある質問（Q&A 3問）
8. 会社概要・実績

各セクションのタイトルは【】で囲み、本文と明確に区別してください。
"""

    response = client.messages.create(
        model="claude-opus-4-7",
        max_tokens=4000,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": prompt}]
    )

    full_text = response.content[0].text
    sections = parse_sections(full_text)
    return sections

def parse_sections(text: str) -> dict:
    """生成テキストをセクション別に分解"""
    import re
    sections = {}
    current_title = "前文"
    current_content = []

    for line in text.split("\n"):
        match = re.match(r"【(.+?)】", line)
        if match:
            sections[current_title] = "\n".join(current_content).strip()
            current_title = match.group(1)
            current_content = [re.sub(r"【.+?】", "", line).strip()]
        else:
            current_content.append(line)

    sections[current_title] = "\n".join(current_content).strip()
    return sections

def build_docx(sections: dict, client_name: str, service: str, output_path: Path):
    """Word文書を組み立て"""
    doc = Document()

    # ページマージン
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # カバーページ
    title = doc.add_heading("", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{service}\nビジネス提案書")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6E)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"提案先：{client_name}").font.size = Pt(14)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(f"{datetime.date.today().strftime('%Y年%m月%d日')}").font.size = Pt(12)

    doc.add_page_break()

    # 各セクション
    for title_text, content in sections.items():
        if title_text == "前文" or not content.strip():
            continue

        heading = doc.add_heading(title_text, level=1)
        heading_run = heading.runs[0] if heading.runs else heading.add_run(title_text)
        heading_run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6E)

        for para_text in content.split("\n"):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.style.font.size = Pt(11)

        doc.add_paragraph()

    # フッター
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.add_run(f"本資料は機密情報を含みます | {client_name} 御中").font.size = Pt(9)

    doc.save(output_path)
    print(f"✅ 提案書を保存: {output_path}")

def main():
    parser = argparse.ArgumentParser(description="AI提案書ジェネレーター")
    parser.add_argument("--client", required=True, help="クライアント企業名")
    parser.add_argument("--service", required=True, help="提案するサービス・製品")
    parser.add_argument("--audience", default="経営者・担当者", help="ターゲット読者")
    parser.add_argument("--budget", default="要相談", help="予算感")
    parser.add_argument("--extra", default="", help="追加情報")
    args = parser.parse_args()

    output_dir = Path(__file__).parent / "samples"
    output_dir.mkdir(exist_ok=True)
    safe_name = args.client.replace(" ", "_").replace("　", "_")
    output_path = output_dir / f"{safe_name}_proposal.docx"

    print(f"🤖 Claude が提案書を生成中... ({args.service} → {args.client})")
    sections = generate_proposal_text(args.client, args.service, args.audience, args.budget, args.extra)

    print("📄 Word文書を組み立て中...")
    build_docx(sections, args.client, args.service, output_path)
    print(f"\n🎉 完成! → {output_path}")

if __name__ == "__main__":
    main()
